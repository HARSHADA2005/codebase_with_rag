"""
Document processing module for the RAG application.
Handles extraction and processing of various document formats.
"""

import io
from pathlib import Path
from typing import List, Dict, Optional
import streamlit as st

# Document processing libraries
import PyPDF2
from docx import Document
import pandas as pd

import config
import utils


class DocumentProcessor:
    """Handles document loading and text extraction."""
    
    def __init__(self):
        self.supported_formats = config.SUPPORTED_FORMATS
    
    @st.cache_data(ttl=config.CACHE_TTL, show_spinner=False)
    def process_file(_self, file_content: bytes, filename: str) -> Dict:
        """
        Process uploaded file and extract text.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dictionary with processed document data
        """
        timer = utils.Timer().start()
        
        file_ext = Path(filename).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                text = _self._extract_pdf(file_content)
            elif file_ext == '.docx':
                text = _self._extract_docx(file_content)
            elif file_ext == '.txt':
                text = _self._extract_txt(file_content)
            elif file_ext == '.csv':
                text = _self._extract_csv(file_content)
            else:
                raise ValueError(f"Unsupported format: {file_ext}")
            
            # Clean the extracted text
            text = utils.clean_text(text)
            
            # Chunk the text
            chunks = utils.chunk_text(text)
            
            timer.stop()
            
            return {
                'filename': filename,
                'text': text,
                'chunks': chunks,
                'num_chunks': len(chunks),
                'num_characters': len(text),
                'processing_time': timer.elapsed(),
                'file_hash': utils.get_file_hash(file_content)
            }
            
        except Exception as e:
            raise Exception(f"Error processing {filename}: {str(e)}")
    
    def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        text = []
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(f"[Page {page_num + 1}]\n{page_text}")
            
            return "\n\n".join(text)
            
        except Exception as e:
            raise Exception(f"PDF extraction error: {str(e)}")
    
    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text.append(" | ".join(row_text))
            
            return "\n\n".join(text)
            
        except Exception as e:
            raise Exception(f"DOCX extraction error: {str(e)}")
    
    def _extract_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file")
            
        except Exception as e:
            raise Exception(f"TXT extraction error: {str(e)}")
    
    def _extract_csv(self, file_content: bytes) -> str:
        """Extract text from CSV file."""
        try:
            csv_file = io.BytesIO(file_content)
            df = pd.read_csv(csv_file)
            
            # Convert dataframe to text representation
            text = [f"CSV Data with {len(df)} rows and {len(df.columns)} columns\n"]
            text.append(f"Columns: {', '.join(df.columns)}\n")
            
            # Add each row as text
            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text.append(row_text)
            
            return "\n".join(text)
            
        except Exception as e:
            raise Exception(f"CSV extraction error: {str(e)}")
    
    def process_multiple_files(self, files: List) -> List[Dict]:
        """
        Process multiple files.
        
        Args:
            files: List of uploaded files
            
        Returns:
            List of processed document dictionaries
        """
        processed_docs = []
        
        for file in files:
            try:
                # Validate file
                is_valid, error_msg = utils.validate_file(file)
                if not is_valid:
                    utils.show_error(f"{file.name}: {error_msg}")
                    continue
                
                # Read file content
                file_content = file.read()
                
                # Process file
                doc_data = self.process_file(file_content, file.name)
                processed_docs.append(doc_data)
                
            except Exception as e:
                utils.show_error(f"Error processing {file.name}: {str(e)}")
        
        return processed_docs


def get_document_stats(doc_data: Dict) -> str:
    """
    Generate statistics string for a processed document.
    
    Args:
        doc_data: Processed document dictionary
        
    Returns:
        Formatted statistics string
    """
    stats = [
        f"📄 **{doc_data['filename']}**",
        f"- Characters: {doc_data['num_characters']:,}",
        f"- Chunks: {doc_data['num_chunks']}",
        f"- Processing time: {doc_data['processing_time']:.2f}s"
    ]
    return "\n".join(stats)
